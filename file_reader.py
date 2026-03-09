# Чтение файлов для передачи агенту Маркус
# Поддерживает: PDF, TXT, MD, CSV, JSON, DOCX, XLSX, изображения (base64)
# Используется в ai_worker.py и window.py

import os
import base64

# Максимум символов текста из файла (чтобы не взорвать контекст)
MAX_TEXT_CHARS = 12_000

SUPPORTED_TEXT_EXTS  = {".txt", ".md", ".py", ".js", ".ts", ".html", ".css",
                         ".json", ".yaml", ".yml", ".xml", ".csv", ".log", ".ini",
                         ".env", ".toml", ".sql", ".bat", ".sh", ".cfg", ".conf"}
SUPPORTED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
SUPPORTED_PDF_EXT    = {".pdf"}
SUPPORTED_DOCX_EXT   = {".docx"}
SUPPORTED_XLSX_EXT   = {".xlsx", ".xls"}


def read_file(path: str) -> dict:
    """
    Читает файл и возвращает словарь:
    {
        "type":     "text" | "image" | "pdf" | "error",
        "name":     "filename.ext",
        "content":  str (текст) или bytes (base64 для изображений),
        "media_type": "image/png" и т.д. (только для изображений),
        "truncated": True/False,
        "size_kb":  float
    }
    """
    if not os.path.exists(path):
        return {"type": "error", "name": path, "content": f"Файл не найден: {path}"}

    name     = os.path.basename(path)
    ext      = os.path.splitext(name)[1].lower()
    size_kb  = os.path.getsize(path) / 1024

    # ── Текстовые файлы ───────────────────────────────────────────────────────
    if ext in SUPPORTED_TEXT_EXTS:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            truncated = len(text) > MAX_TEXT_CHARS
            return {
                "type":      "text",
                "name":      name,
                "content":   text[:MAX_TEXT_CHARS],
                "truncated": truncated,
                "size_kb":   size_kb,
            }
        except Exception as e:
            return {"type": "error", "name": name, "content": str(e)}

    # ── PDF ───────────────────────────────────────────────────────────────────
    if ext in SUPPORTED_PDF_EXT:
        return _read_pdf(path, name, size_kb)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if ext in SUPPORTED_DOCX_EXT:
        return _read_docx(path, name, size_kb)

    # ── XLSX / XLS ────────────────────────────────────────────────────────────
    if ext in SUPPORTED_XLSX_EXT:
        return _read_xlsx(path, name, size_kb)

    # ── Изображения ───────────────────────────────────────────────────────────
    if ext in SUPPORTED_IMAGE_EXTS:
        return _read_image(path, name, ext, size_kb)

    # ── Неизвестный формат — пробуем как текст ────────────────────────────────
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type":      "text",
            "name":      name,
            "content":   text[:MAX_TEXT_CHARS],
            "truncated": truncated,
            "size_kb":   size_kb,
        }
    except Exception:
        return {
            "type":    "error",
            "name":    name,
            "content": f"Формат {ext} не поддерживается.",
        }


def _read_pdf(path, name, size_kb):
    # Сначала пробуем PyMuPDF (fitz), потом pdfplumber, потом pypdf
    try:
        import fitz  # PyMuPDF
        doc  = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        pass

    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        text = "\n".join(text_parts)
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        pass

    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text   = "\n".join(p.extract_text() or "" for p in reader.pages)
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        pass

    return {"type": "error", "name": name,
            "content": "Для PDF установи: pip install PyMuPDF или pdfplumber"}


def _read_docx(path, name, size_kb):
    # Способ 1: python-docx (если установлен)
    try:
        import docx
        doc   = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        truncated = len(text) > MAX_TEXT_CHARS
        print(f"[FILE] DOCX прочитан через python-docx: {len(text)} символов")
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        pass  # Нет python-docx — пробуем zipfile
    except Exception as e:
        return {"type": "error", "name": name, "content": f"python-docx ошибка: {e}"}

    # Способ 2: zipfile — docx это ZIP с XML, работает без зависимостей
    try:
        import zipfile
        import re
        with zipfile.ZipFile(path, "r") as z:
            with z.open("word/document.xml") as f:
                xml = f.read().decode("utf-8", errors="replace")
        # Переводы строк на месте тегов абзацев и переносов
        text = re.sub(r"<w:br[^/]*/?>", "\n", xml)
        text = re.sub(r"</w:p>", "\n", text)
        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        truncated = len(text) > MAX_TEXT_CHARS
        print(f"[FILE] DOCX прочитан через zipfile (без python-docx): {len(text)} символов")
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except Exception as e:
        return {
            "type": "error", "name": name,
            "content": f"Не удалось прочитать DOCX: {e}\nУстанови: pip install python-docx --break-system-packages",
        }


def _read_xlsx(path, name, size_kb):
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(path, read_only=True, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"=== Лист: {sheet} ===")
            for row in ws.iter_rows(values_only=True):
                row_str = "\t".join(str(c) if c is not None else "" for c in row)
                if row_str.strip():
                    lines.append(row_str)
        text = "\n".join(lines)
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        pass

    try:
        import pandas as pd
        xf   = pd.ExcelFile(path)
        lines = []
        for sheet in xf.sheet_names:
            df = xf.parse(sheet)
            lines.append(f"=== Лист: {sheet} ===")
            lines.append(df.to_string(index=False))
        text = "\n".join(lines)
        truncated = len(text) > MAX_TEXT_CHARS
        return {
            "type": "text", "name": name,
            "content": text[:MAX_TEXT_CHARS],
            "truncated": truncated, "size_kb": size_kb,
        }
    except ImportError:
        return {"type": "error", "name": name,
                "content": "Для XLSX установи: pip install openpyxl"}
    except Exception as e:
        return {"type": "error", "name": name, "content": str(e)}


def _read_image(path, name, ext, size_kb):
    media_map = {
        ".png":  "image/png",
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif":  "image/gif",
    }
    try:
        with open(path, "rb") as f:
            raw = f.read()
        b64 = base64.standard_b64encode(raw).decode("utf-8")
        return {
            "type":       "image",
            "name":       name,
            "content":    b64,
            "media_type": media_map.get(ext, "image/png"),
            "truncated":  False,
            "size_kb":    size_kb,
        }
    except Exception as e:
        return {"type": "error", "name": name, "content": str(e)}


def format_for_prompt(file_info: dict) -> str:
    """
    Форматирует файл как текст для вставки в промпт.
    Используется когда модель не поддерживает vision.
    """
    t    = file_info["type"]
    name = file_info["name"]
    if t == "error":
        return f"[ФАЙЛ {name}: ОШИБКА — {file_info['content']}]"
    if t == "image":
        return f"[ФАЙЛ {name}: изображение, {file_info['size_kb']:.1f} КБ — передано как base64]"
    trunc = " [ТЕКСТ ОБРЕЗАН]" if file_info.get("truncated") else ""
    return (
        f"\n--- ФАЙЛ: {name} ({file_info['size_kb']:.1f} КБ){trunc} ---\n"
        f"{file_info['content']}\n"
        f"--- КОНЕЦ ФАЙЛА: {name} ---\n"
    )